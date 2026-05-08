"""
Permohonan Draft Storage
========================
Per-project storage and DOCX export for generated petition drafts.
"""

from __future__ import annotations

import json
import os
import re
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt

from .project_store import PROJECTS_DIR, get_project


def _now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _safe_id(value: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_-]", "", value or "")


def _project_dir(project_id: str, projects_dir: str = PROJECTS_DIR) -> Path:
    return Path(projects_dir) / _safe_id(project_id)


def _drafts_dir(project_id: str, projects_dir: str = PROJECTS_DIR) -> Path:
    return _project_dir(project_id, projects_dir) / "drafts"


def _make_title(user_input: Dict[str, Any], mode: str, fallback: str = "") -> str:
    pemohon = str(user_input.get("nama_pemohon") or "").strip()
    uu = str(user_input.get("uu_diuji") or "").strip()
    if pemohon and uu:
        return f"Permohonan {pemohon} - {uu}"
    if uu:
        return f"Permohonan Pengujian {uu}"
    if pemohon:
        return f"Permohonan {pemohon}"
    if fallback:
        return fallback
    return "Perbaikan Draft Permohonan" if mode == "improve_existing_draft" else "Draft Permohonan Baru"


def _add_paragraph(document: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        document.add_paragraph("")
        return

    is_heading = (
        stripped.isupper()
        or bool(re.match(r"^(?:[IVXLCDM]+\.|[A-Z]\.|BAB\s+|CATATAN\s+DRAFTER)", stripped))
    )
    paragraph = document.add_paragraph()
    run = paragraph.add_run(stripped)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if is_heading:
        run.bold = True
        paragraph.space_before = Pt(8)
        paragraph.space_after = Pt(4)
    else:
        paragraph.paragraph_format.first_line_indent = Pt(28)
        paragraph.paragraph_format.space_after = Pt(3)


def write_docx(draft_text: str, path: Path, title: str = "Draft Permohonan") -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    document.add_heading(title, level=1)
    for block in re.split(r"\n+", draft_text.strip()):
        _add_paragraph(document, block)
    document.save(path)


def save_permohonan_draft(
    project_id: str,
    mode: str,
    user_input: Dict[str, Any],
    draft_text: str,
    uploaded_draft: Optional[Dict[str, Any]] = None,
    sources: Optional[Dict[str, Any]] = None,
    projects_dir: str = PROJECTS_DIR,
) -> Optional[Dict[str, Any]]:
    if not get_project(project_id):
        return None
    if not draft_text.strip():
        return None

    draft_id = uuid.uuid4().hex[:10]
    timestamp = _now_iso()
    title = _make_title(user_input, mode, uploaded_draft.get("filename", "") if uploaded_draft else "")
    drafts_dir = _drafts_dir(project_id, projects_dir)
    drafts_dir.mkdir(parents=True, exist_ok=True)

    txt_name = f"draft_{draft_id}.txt"
    docx_name = f"draft_{draft_id}.docx"
    meta_name = f"draft_{draft_id}.json"
    txt_path = drafts_dir / txt_name
    docx_path = drafts_dir / docx_name
    meta_path = drafts_dir / meta_name

    txt_path.write_text(draft_text, encoding="utf-8")
    write_docx(draft_text, docx_path, title)

    metadata = {
        "id": draft_id,
        "title": title,
        "mode": mode,
        "timestamp": timestamp,
        "user_input": user_input,
        "uploaded_draft": {
            "filename": (uploaded_draft or {}).get("filename", ""),
            "text_chars": len((uploaded_draft or {}).get("raw_text", "") or ""),
        },
        "sources": sources or {},
        "txt_filename": txt_name,
        "docx_filename": docx_name,
        "draft_excerpt": draft_text[:500],
        "draft_chars": len(draft_text),
    }
    meta_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
    return metadata


def list_permohonan_drafts(project_id: str, projects_dir: str = PROJECTS_DIR) -> List[Dict[str, Any]]:
    drafts_dir = _drafts_dir(project_id, projects_dir)
    if not drafts_dir.is_dir():
        return []

    drafts: List[Dict[str, Any]] = []
    for path in drafts_dir.glob("draft_*.json"):
        try:
            drafts.append(json.loads(path.read_text(encoding="utf-8")))
        except Exception:
            continue
    drafts.sort(key=lambda item: item.get("timestamp", ""), reverse=True)
    return drafts


def get_permohonan_draft_docx_path(
    project_id: str,
    draft_id: str,
    projects_dir: str = PROJECTS_DIR,
) -> Optional[Path]:
    safe_draft_id = _safe_id(draft_id)
    drafts_dir = _drafts_dir(project_id, projects_dir).resolve()
    meta_path = drafts_dir / f"draft_{safe_draft_id}.json"
    if not meta_path.exists():
        return None

    try:
        metadata = json.loads(meta_path.read_text(encoding="utf-8"))
    except Exception:
        return None

    docx_path = (drafts_dir / metadata.get("docx_filename", "")).resolve()
    if not str(docx_path).startswith(str(drafts_dir)):
        return None
    if not docx_path.exists():
        return None
    return docx_path
